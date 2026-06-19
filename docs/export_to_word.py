"""Export INF612 report markdown to Word (.docx)."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

DOCS_DIR = Path(__file__).resolve().parent
INPUT_MD = DOCS_DIR / "INF612_Assignment2_Report.md"
OUTPUT_DOCX = DOCS_DIR / "INF612_Assignment2_Report.docx"


def add_formatted_text(paragraph, text: str) -> None:
    """Add text with basic **bold** and *italic* handling."""
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
        else:
            paragraph.add_run(part)


def parse_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\|?[\s\-:|]+\|?$", line.strip()))


def export_markdown_to_docx(md_path: Path, docx_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    in_code_block = False
    code_buffer: list[str] = []
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code_block:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_buffer))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        if stripped in ("---", "***"):
            doc.add_paragraph()
            i += 1
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=0)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=1)
            i += 1
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
            i += 1
            continue
        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:].strip(), level=3)
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            headers = parse_table_row(stripped)
            i += 2
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(parse_table_row(lines[i]))
                i += 1
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = "Table Grid"
            for col, header in enumerate(headers):
                table.rows[0].cells[col].text = header
            for r, row in enumerate(rows, start=1):
                for col, cell in enumerate(row):
                    if col < len(table.rows[r].cells):
                        table.rows[r].cells[col].text = cell
            doc.add_paragraph()
            continue

        if stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            add_formatted_text(p, stripped[2:])
            i += 1
            continue

        if stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            add_formatted_text(p, stripped[2:])
            i += 1
            continue

        numbered = re.match(r"^(\d+)\.\s+(.*)", stripped)
        if numbered:
            p = doc.add_paragraph(style="List Number")
            add_formatted_text(p, numbered.group(2))
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped.startswith("**Keywords:**"):
            p = doc.add_paragraph()
            add_formatted_text(p, stripped)
            i += 1
            continue

        p = doc.add_paragraph()
        add_formatted_text(p, stripped)
        i += 1

    doc.save(docx_path)
    print(f"Exported: {docx_path}")


if __name__ == "__main__":
    export_markdown_to_docx(INPUT_MD, OUTPUT_DOCX)
